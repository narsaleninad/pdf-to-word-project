from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_word(lines, output_path):
    doc = Document()

    for line in lines:
        if line == "__PAGE_BREAK__":
            doc.add_page_break()
            continue

        if not line:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        r = p.add_run(line)

        # Heading logic
        if line.isupper() or line.endswith(":"):
            r.bold = True
            r.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
        else:
            r.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
